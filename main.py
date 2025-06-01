"""
Improved table image to docx converter implementation
Main improvements:
1. Smarter line detection algorithm
2. When vertical lines are insufficient, attempt to infer table structure from image content
3. Support for incomplete table lines
4. Better error handling and fallback mechanisms

Steps:
1. Load table image
2. Remove text
3. Detect all possible table lines (Initial detection)
3b. Refine lines using Deep Learning model
4. Merge & filter lines
5. Intelligently supplement missing lines
6. Build 2D table data structure
7. Generate corresponding docx table
8. Write to docx file, complete conversion
"""

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
from typing import List, Tuple, Dict, Optional
import os
import numpy as np 
import cv2
import tensorflow as tf 
import pytesseract 
from recover_outline import recover_outline_AC, recover_outline_UN

# PATHS 
input_folder = "dataset/application"
output_docx_folder = "dataset/output_docx_app"
visualization_folder = "dataset/visualizations_app"
denoised_folder = "dataset/denoised" 
model_weights_path = 'model/m.weights.h5'
ae_model_path = 'model/AutoEncoder.pth'
un_model_path = 'model/UNet.pth'

denoise_model = 'AutoEncoder'  # Options: 'AutoEncoder', 'UNet'


class ImprovedTableProcessor:
    def __init__(self, model_weights_path=model_weights_path): 
        self.image = None
        self.original_image = None
        self.text_removed_image = None
        self.denoised_image = None
        self.horizontal_lines = []
        self.vertical_lines = []
        self.table_structure = []
        self.extracted_text = {}
        self.image_height = 0
        self.image_width = 0
        self.dl_model = self._load_dl_model(model_weights_path)

    def _load_dl_model(self, weights_path: str) -> Optional[tf.keras.Model]:
        """
        Loads the pre-trained Keras LSTM model for line refinement.
        """
        try:
            model = tf.keras.Sequential([
                tf.keras.layers.LSTM(128, return_sequences=True, input_shape=(None, 4)),
                tf.keras.layers.LSTM(128, return_sequences=True),
                tf.keras.layers.Dense(4)
            ])
            if os.path.exists(weights_path):
                print(f"Loading DL model weights from {weights_path}")
                model.load_weights(weights_path)
                return model
            else:
                print(f"Warning: DL model weights not found at {weights_path}. DL refinement will be skipped.")
                return None
        except Exception as e:
            print(f"Error loading DL model: {e}. DL refinement will be skipped.")
            return None
    
    def step1_load_image(self, image_path: str) -> np.ndarray:
        """
        Step 1: Load table image
        """
        # read image
        self.original_image = cv2.imread(image_path)
        if self.original_image is None:
            raise ValueError(f"Cannot read image: {image_path}")
            
        # Convert to grayscale
        self.image = cv2.cvtColor(self.original_image, cv2.COLOR_BGR2GRAY)
        self.image_height, self.image_width = self.image.shape
        
        return self.image
    
    def step2_remove_text(self) -> np.ndarray:
        """
        Step 2: Remove text (improved version - more sensitive to thin lines)
        """
        if self.image is None:
            raise ValueError("Please load image first")
        
        binary_image = cv2.adaptiveThreshold(
            self.image, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, 
            cv2.THRESH_BINARY_INV, 11, 2  # Block size 11, C 2
        )
        
        # Create multiple sized structuring elements to detect lines of different thickness
        horizontal_kernels = [
            cv2.getStructuringElement(cv2.MORPH_RECT, (8, 1)), 
            cv2.getStructuringElement(cv2.MORPH_RECT, (10, 1)),
            cv2.getStructuringElement(cv2.MORPH_RECT, (20, 1)),
            cv2.getStructuringElement(cv2.MORPH_RECT, (40, 1)),
            cv2.getStructuringElement(cv2.MORPH_RECT, (60, 1))
        ]
        
        vertical_kernels = [
            cv2.getStructuringElement(cv2.MORPH_RECT, (1, 8)), 
            cv2.getStructuringElement(cv2.MORPH_RECT, (1, 10)),
            cv2.getStructuringElement(cv2.MORPH_RECT, (1, 20)),
            cv2.getStructuringElement(cv2.MORPH_RECT, (1, 40)),
            cv2.getStructuringElement(cv2.MORPH_RECT, (1, 60))
        ]
        
        # Extract lines using multiple kernels on the binary image
        horizontal_lines_combined = np.zeros_like(binary_image)
        for kernel in horizontal_kernels:
            h_lines = cv2.morphologyEx(binary_image, cv2.MORPH_OPEN, kernel)
            horizontal_lines_combined = cv2.bitwise_or(horizontal_lines_combined, h_lines)
        
        vertical_lines_combined = np.zeros_like(binary_image)
        for kernel in vertical_kernels:
            v_lines = cv2.morphologyEx(binary_image, cv2.MORPH_OPEN, kernel)
            vertical_lines_combined = cv2.bitwise_or(vertical_lines_combined, v_lines)
        
        # Combine horizontal and vertical lines
        self.text_removed_image = cv2.bitwise_or(horizontal_lines_combined, vertical_lines_combined)
        
        # Filter out small connected components likely to be text remnants or noise.
        # Increased min_area for better noise removal.
        num_labels, labels, stats, _ = cv2.connectedComponentsWithStats(self.text_removed_image, connectivity=8)
        
        min_area = 10  
        
        filtered_image = np.zeros_like(self.text_removed_image)
        for i in range(1, num_labels):  # Skip background label 0
            if stats[i, cv2.CC_STAT_AREA] >= min_area:
                filtered_image[labels == i] = 255
        
        self.text_removed_image = filtered_image
        # Reverse black and white
        self.text_removed_image = cv2.bitwise_not(self.text_removed_image)
        
        return self.text_removed_image
    
    def step2_5_denoise_image(self) -> np.ndarray:
        """
        Step 2.5: Denoise the image after text removal
        This step cleans up noise and artifacts while preserving table structure
        """
        if self.text_removed_image is None:
            raise ValueError("Please execute text removal step first")
        if self.original_image is None:
            raise ValueError("Original image is not loaded. Cannot create full comparison.")

        # Store original dimensions
        original_height, original_width = self.text_removed_image.shape
        
        target_height = ((original_height + 31) // 32) * 32  # Round up to nearest multiple of 32
        target_width = ((original_width + 31) // 32) * 32   # Round up to nearest multiple of 32
        
        # Resize image for denoising if dimensions don't match requirements
        if original_height != target_height or original_width != target_width:
            resized_image = cv2.resize(self.text_removed_image, (target_width, target_height), 
                                     interpolation=cv2.INTER_LINEAR)
        else:
            resized_image = self.text_removed_image

        # Apply denoising on the resized image
        try:
            if denoise_model == 'AutoEncoder':
                denoised_resized = recover_outline_AC(resized_image, ae_model_path)
            elif denoise_model == 'UNet':
                denoised_resized = recover_outline_UN(resized_image, un_model_path)
            else:
                # Fallback to original image if model not recognized
                denoised_resized = resized_image
        except Exception as e:
            print(f"Warning: Denoising failed with error: {e}. Using original text-removed image.")
            denoised_resized = resized_image

        # Resize back to original dimensions if we resized earlier
        if original_height != target_height or original_width != target_width:
            self.denoised_image = cv2.resize(denoised_resized, (original_width, original_height), 
                                           interpolation=cv2.INTER_LINEAR)
        else:
            self.denoised_image = denoised_resized

        return self.denoised_image 
    
    def step3_detect_table_lines(self) -> Tuple[List, List]:
        """
        Step 3: Detect all possible table lines (improved version - preserve original line lengths, no extension)
        """
        image_for_detection = self.denoised_image

        # Invert the image for detection
        detection_input_image = cv2.bitwise_not(image_for_detection)
        
        self.horizontal_lines = []
        self.vertical_lines = []
        
        # Define boundary area (exclude lines at image edges)
        border_margin = 5 
        
        # Method 1: Hough transform detection (preserve original line lengths)
        # Use detection_input_image (white lines on black background)
        edges = cv2.Canny(detection_input_image, 20, 80, apertureSize=3)
        lines = cv2.HoughLinesP(edges, 1, np.pi/180, threshold=15, minLineLength=10, maxLineGap=5)
        
        if lines is not None:
            for line in lines:
                x1, y1, x2, y2 = line[0]
                
                # Exclude lines near image boundaries
                if (x1 < border_margin or x2 < border_margin or 
                    x1 > self.image_width - border_margin or x2 > self.image_width - border_margin or
                    y1 < border_margin or y2 < border_margin or 
                    y1 > self.image_height - border_margin or y2 > self.image_height - border_margin):
                    continue
                
                # Calculate line angle and length
                angle = np.arctan2(y2 - y1, x2 - x1) * 180 / np.pi
                length = np.sqrt((x2 - x1)**2 + (y2 - y1)**2)
                
                if length < 10:
                    continue
                
                # Classify horizontal and vertical lines - PRESERVE ORIGINAL COORDINATES
                angle_tolerance = 10 
                if abs(angle) < angle_tolerance or abs(angle) > (180 - angle_tolerance):  # Horizontal lines
                    self.horizontal_lines.append((x1, y1, x2, y2))
                elif abs(angle - 90) < angle_tolerance or abs(angle + 90) < angle_tolerance:  # Vertical lines
                    self.vertical_lines.append((x1, y1, x2, y2))
        
        # Method 2: Direct analysis from morphological operation results (modified)
        # Pass detection_input_image (white lines on black background)
        self._detect_lines_from_projections_preserve_length(detection_input_image)
        
        return self.horizontal_lines, self.vertical_lines
    
    def _detect_lines_from_projections_preserve_length(self, image_to_process: np.ndarray):
        """
        Detect lines using projection analysis but preserve actual line lengths
        Uses the provided image_to_process (expected to be white lines on black background)
        """
        border_margin = 10
        
        # Horizontal projection - detect horizontal lines with actual lengths
        for y in range(border_margin, self.image_height - border_margin):
            row = image_to_process[y, :]
            
            # Find continuous segments of pixels that form lines
            line_segments = []
            start_x = None
            
            for x in range(len(row)):
                if row[x] > 50:  # Pixel is part of a line
                    if start_x is None:
                        start_x = x
                else:  # End of line segment
                    if start_x is not None:
                        segment_length = x - start_x
                        if segment_length > self.image_width * 0.1:  # Minimum meaningful length
                            line_segments.append((start_x, y, x-1, y))
                        start_x = None
            
            # Handle case where line extends to end of row
            if start_x is not None:
                segment_length = len(row) - start_x
                if segment_length > self.image_width * 0.1:
                    line_segments.append((start_x, y, len(row)-1, y))
            
            # Add detected line segments
            for segment in line_segments:
                x1, y1, x2, y2 = segment
                if x1 > border_margin and x2 < self.image_width - border_margin:
                    self.horizontal_lines.append((x1, y1, x2, y2))
        
        # Vertical projection - detect vertical lines with actual lengths
        for x in range(border_margin, self.image_width - border_margin):
            col = image_to_process[:, x]
            
            # Find continuous segments of pixels that form lines
            line_segments = []
            start_y = None
            
            for y in range(len(col)):
                if col[y] > 50:  # Pixel is part of a line
                    if start_y is None:
                        start_y = y
                else:  # End of line segment
                    if start_y is not None:
                        segment_length = y - start_y
                        if segment_length > self.image_height * 0.1:  # Minimum meaningful length
                            line_segments.append((x, start_y, x, y-1))
                        start_y = None
            
            # Handle case where line extends to end of column
            if start_y is not None:
                segment_length = len(col) - start_y
                if segment_length > self.image_height * 0.1:
                    line_segments.append((x, start_y, x, len(col)-1))
            
            # Add detected line segments
            for segment in line_segments:
                x1, y1, x2, y2 = segment
                if y1 > border_margin and y2 < self.image_height - border_margin:
                    self.vertical_lines.append((x1, y1, x2, y2))

    def step4_merge_and_filter_lines(self) -> Tuple[List, List]:
        """
        Step 4: Merge & filter lines (modified to preserve partial line lengths)
        """
        def merge_lines_preserve_length(lines, is_horizontal=True, distance_threshold=20):
            if not lines:
                return []
            
            merged_lines = []
            lines = sorted(lines)
            
            for line in lines:
                x1, y1, x2, y2 = line
                
                # Ensure consistent line direction
                if is_horizontal and x1 > x2:
                    x1, y1, x2, y2 = x2, y2, x1, y1
                elif not is_horizontal and y1 > y2:
                    x1, y1, x2, y2 = x2, y2, x1, y1
                
                merged = False
                for i, existing_line in enumerate(merged_lines):
                    ex1, ey1, ex2, ey2 = existing_line
                    
                    if is_horizontal:
                        # Merge horizontal lines only if they overlap significantly
                        if abs(y1 - ey1) < distance_threshold:
                            overlap_start = max(x1, ex1)
                            overlap_end = min(x2, ex2)
                            if overlap_end > overlap_start:  # Lines actually overlap
                                # Merge by extending the union, not just taking max range
                                new_x1 = min(x1, ex1)
                                new_x2 = max(x2, ex2)
                                new_y = (y1 + ey1) // 2
                                merged_lines[i] = (new_x1, new_y, new_x2, new_y)
                                merged = True
                                break
                    else:
                        # Merge vertical lines only if they overlap significantly
                        if abs(x1 - ex1) < distance_threshold:
                            overlap_start = max(y1, ey1)
                            overlap_end = min(y2, ey2)
                            if overlap_end > overlap_start:  # Lines actually overlap
                                # Merge by extending the union
                                new_y1 = min(y1, ey1)
                                new_y2 = max(y2, ey2)
                                new_x = (x1 + ex1) // 2
                                merged_lines[i] = (new_x, new_y1, new_x, new_y2)
                                merged = True
                                break
                
                if not merged:
                    merged_lines.append((x1, y1, x2, y2))
            
            return merged_lines
        
        # Merge lines while preserving lengths
        self.horizontal_lines = merge_lines_preserve_length(self.horizontal_lines, True)
        self.vertical_lines = merge_lines_preserve_length(self.vertical_lines, False)
        
        # Filter lines that are too short, but keep partial lines
        min_h_length = max(10, self.image_width * 0.05)  
        min_v_length = max(10, self.image_height * 0.05)
        
        border_margin = 8
        
        filtered_h_lines = []
        for line in self.horizontal_lines:
            x1, y1, x2, y2 = line
            length = abs(x2 - x1)
            if (y1 > border_margin and y1 < self.image_height - border_margin and 
                length >= min_h_length):
                filtered_h_lines.append(line)
        
        filtered_v_lines = []
        for line in self.vertical_lines:
            x1, y1, x2, y2 = line
            length = abs(y2 - y1)
            if (x1 > border_margin and x1 < self.image_width - border_margin and 
                length >= min_v_length):
                filtered_v_lines.append(line)
        
        self.horizontal_lines = filtered_h_lines
        self.vertical_lines = filtered_v_lines
        
        return self.horizontal_lines, self.vertical_lines
    
    def step5_smart_line_completion(self) -> Tuple[List, List]:
        """
        Step 5: Intelligently supplement missing lines
        """
        # If too few vertical lines, try to infer column boundaries from text distribution
        if len(self.vertical_lines) < 2:
            self._infer_vertical_lines_from_text()
        
        # If too few horizontal lines, try to infer row boundaries from text distribution
        if len(self.horizontal_lines) < 2:
            self._infer_horizontal_lines_from_text()
        
        # Ensure at least boundary lines exist
        self._ensure_boundary_lines()
        
        return self.horizontal_lines, self.vertical_lines

    def _infer_vertical_lines_from_text(self):
        """
        Infer vertical line positions from text distribution with improved OCR
        """
        try:
            ocr_data = self._try_multiple_ocr_configs(self.image)
            
            # Collect valid text regions with improved validation
            text_regions = []
            valid_texts_found = []
            
            for i in range(len(ocr_data.get('text', []))):
                confidence = int(ocr_data['conf'][i]) if i < len(ocr_data['conf']) else 0
                text = ocr_data['text'][i].strip() if i < len(ocr_data['text']) else ""
                
                # More lenient initial filtering
                if confidence > 30 and len(text) > 0:
                    # Clean and validate text
                    clean_text = ''.join(char for char in text if char.isprintable()).strip()
                    
                    # Accept text that has any alphanumeric characters or common symbols
                    if (len(clean_text) > 0 and 
                        (any(char.isalnum() for char in clean_text) or 
                         any(char in clean_text for char in ['$', '%', '.', ',', '-', '+']))):
                        
                        x = ocr_data['left'][i] if i < len(ocr_data['left']) else 0
                        y = ocr_data['top'][i] if i < len(ocr_data['top']) else 0
                        w = ocr_data['width'][i] if i < len(ocr_data['width']) else 0
                        h = ocr_data['height'][i] if i < len(ocr_data['height']) else 0
                        
                        if w > 3 and h > 3:  # Minimum size requirement
                            text_regions.append((x, y, x+w, y+h))
                            valid_texts_found.append(clean_text)
            
            # try edge-based text detection
            if len(text_regions) < 2:
                text_regions = self._detect_text_regions_by_edges()
            
            if len(text_regions) >= 2:
                # Analyze x-coordinate distribution
                left_edges = [region[0] for region in text_regions]
                right_edges = [region[2] for region in text_regions]
                
                # Use clustering analysis to find possible column boundaries
                all_x_coords = sorted(set(left_edges + right_edges))
                
                # Simple clustering: find clusters of x-coordinates
                clusters = []
                current_cluster = [all_x_coords[0]]
                
                for x in all_x_coords[1:]:
                    if x - current_cluster[-1] < 40:  # Increased tolerance
                        current_cluster.append(x)
                    else:
                        clusters.append(current_cluster)
                        current_cluster = [x]
                clusters.append(current_cluster)
                
                if len(clusters) >= 2:
                    border_margin = 15
                    for cluster in clusters:
                        x = int(np.mean(cluster))
                        if border_margin <= x <= self.image_width - border_margin:
                            y1 = border_margin
                            y2 = self.image_height - border_margin
                            self.vertical_lines.append((x, y1, x, y2))
                else:
                    self._add_geometric_vertical_lines()
            else:
                self._add_geometric_vertical_lines()
                
        except Exception as e:
            self._add_geometric_vertical_lines()

    def _detect_text_regions_by_edges(self):
        """
        Detect text regions using edge detection when OCR fails
        """
        try:
            # Edge detection
            edges = cv2.Canny(self.image, 50, 150)
            
            # Find contours
            contours, _ = cv2.findContours(edges, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
            
            text_regions = []
            for contour in contours:
                x, y, w, h = cv2.boundingRect(contour)
                
                # Filter by size - likely text regions
                if (10 < w < self.image_width * 0.3 and 
                    8 < h < self.image_height * 0.1 and
                    w > h):  # Text is usually wider than tall
                    
                    # Check if region has sufficient content
                    roi = self.image[y:y+h, x:x+w]
                    if roi.size > 0:
                        std_dev = np.std(roi)
                        if std_dev > 20:  # Has variation suggesting text
                            text_regions.append((x, y, x+w, y+h))
            
            return text_regions
            
        except Exception as e:
            return []

    def _add_geometric_vertical_lines(self):
        """
        Add vertical lines using geometric division when text analysis fails
        """
        if len(self.vertical_lines) == 0:
            border_margin = 15
            # Try to create reasonable column divisions
            effective_width = self.image_width - 2 * border_margin
            
            # Create 2-3 columns based on image aspect ratio
            if self.image_width > self.image_height * 1.5:  # Wide image, more columns
                num_cols = 3
            else:
                num_cols = 2
            
            for i in range(1, num_cols):
                x = border_margin + (effective_width * i // num_cols)
                y1 = border_margin
                y2 = self.image_height - border_margin
                self.vertical_lines.append((x, y1, x, y2))

    def _infer_horizontal_lines_from_text(self):
        """
        Infer horizontal line positions from text distribution with improved OCR
        """
        try:
            # Try multiple OCR approaches
            ocr_data = self._try_multiple_ocr_configs(self.image)
            
            # Collect valid text regions with improved validation
            text_regions = []
            valid_texts_found = []
            
            for i in range(len(ocr_data.get('text', []))):
                confidence = int(ocr_data['conf'][i]) if i < len(ocr_data['conf']) else 0
                text = ocr_data['text'][i].strip() if i < len(ocr_data['text']) else ""
                
                # More lenient initial filtering
                if confidence > 30 and len(text) > 0:
                    clean_text = ''.join(char for char in text if char.isprintable()).strip()
                    
                    if (len(clean_text) > 0 and 
                        (any(char.isalnum() for char in clean_text) or 
                         any(char in clean_text for char in ['$', '%', '.', ',', '-', '+']))):
                        
                        x = ocr_data['left'][i] if i < len(ocr_data['left']) else 0
                        y = ocr_data['top'][i] if i < len(ocr_data['top']) else 0
                        w = ocr_data['width'][i] if i < len(ocr_data['width']) else 0
                        h = ocr_data['height'][i] if i < len(ocr_data['height']) else 0
                        
                        if w > 3 and h > 3:
                            text_regions.append((x, y, x+w, y+h))
                            valid_texts_found.append(clean_text)
            
            # If still no good OCR results, try edge-based detection
            if len(text_regions) < 2:
                text_regions = self._detect_text_regions_by_edges()
            
            if len(text_regions) >= 2:
                # Analyze y-coordinate distribution
                top_edges = [region[1] for region in text_regions]
                bottom_edges = [region[3] for region in text_regions]
                
                # Use clustering analysis to find possible row boundaries
                all_y_coords = sorted(set(top_edges + bottom_edges))
                
                # Simple clustering: find clusters of y-coordinates
                clusters = []
                current_cluster = [all_y_coords[0]]
                
                for y in all_y_coords[1:]:
                    if y - current_cluster[-1] < 30:  # Increased tolerance
                        current_cluster.append(y)
                    else:
                        clusters.append(current_cluster)
                        current_cluster = [y]
                clusters.append(current_cluster)
                
                if len(clusters) >= 2:
                    border_margin = 15
                    for cluster in clusters:
                        y = int(np.mean(cluster))
                        if border_margin <= y <= self.image_height - border_margin:
                            x1 = border_margin
                            x2 = self.image_width - border_margin
                            self.horizontal_lines.append((x1, y, x2, y))
                else:
                    self._add_geometric_horizontal_lines()
            else:
                self._add_geometric_horizontal_lines()
                
        except Exception as e:
            self._add_geometric_horizontal_lines()

    def _add_geometric_horizontal_lines(self):
        """
        Add horizontal lines using geometric division when text analysis fails
        """
        if len(self.horizontal_lines) == 0:
            border_margin = 15
            effective_height = self.image_height - 2 * border_margin
            
            # Create 2-4 rows based on image aspect ratio
            if self.image_height > self.image_width * 0.8:  # Tall image, more rows
                num_rows = 4
            else:
                num_rows = 2
            
            for i in range(1, num_rows):
                y = border_margin + (effective_height * i // num_rows)
                x1 = border_margin
                x2 = self.image_width - border_margin
                self.horizontal_lines.append((x1, y, x2, y))

    def _ensure_boundary_lines(self):
        """
        Ensure sufficient boundary lines exist, but avoid image boundaries
        """
        border_margin = 15  # Distance to avoid image boundaries
        
        # Check if boundary lines need to be added
        # Only add when line count is severely insufficient
        if len(self.horizontal_lines) < 2:
            # Add internal boundary lines (not image boundaries)
            top_y = border_margin
            bottom_y = self.image_height - border_margin
            x1 = border_margin
            x2 = self.image_width - border_margin
            
            self.horizontal_lines.append((x1, top_y, x2, top_y))
            self.horizontal_lines.append((x1, bottom_y, x2, bottom_y))
        
        if len(self.vertical_lines) < 2:
            # Add internal boundary lines (not image boundaries)
            left_x = border_margin
            right_x = self.image_width - border_margin
            y1 = border_margin
            y2 = self.image_height - border_margin
            
            self.vertical_lines.append((left_x, y1, left_x, y2))
            self.vertical_lines.append((right_x, y1, right_x, y2))

    def _preprocess_image_for_ocr(self, image):
        """
        Preprocess image to improve OCR accuracy
        """
        # Resize image if too small (OCR works better on larger images)
        height, width = image.shape
        if width < 300 or height < 300:
            scale_factor = max(2, 300 / min(width, height))
            new_width = int(width * scale_factor)
            new_height = int(height * scale_factor)
            image = cv2.resize(image, (new_width, new_height), interpolation=cv2.INTER_CUBIC)
        
        # Enhance contrast
        clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8,8))
        image = clahe.apply(image)
        
        # Gaussian blur to reduce noise
        image = cv2.GaussianBlur(image, (1, 1), 0)
        
        # Sharpen image
        kernel = np.array([[-1,-1,-1], [-1,9,-1], [-1,-1,-1]])
        image = cv2.filter2D(image, -1, kernel)
        
        # Ensure good contrast
        image = cv2.convertScaleAbs(image, alpha=1.3, beta=10)
        
        return image

    def _try_multiple_ocr_configs(self, image):
        """
        Try multiple OCR configurations to improve text detection
        """
        # Preprocess image
        processed_image = self._preprocess_image_for_ocr(image)
        
        # Multiple OCR configurations
        configs = [
            '--psm 6 --oem 3',  # Uniform block of text
            '--psm 8 --oem 3',  # Single word
            '--psm 7 --oem 3',  # Single text line
            '--psm 13 --oem 3', # Raw line, treat as single text line
            '--psm 6 --oem 3 -c tessedit_char_whitelist=0123456789ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz .,%-$()[]',
        ]
        
        best_result = None
        best_confidence = 0
        
        for config in configs:
            try:
                # Try with original processed image
                data = pytesseract.image_to_data(processed_image, config=config, 
                                               output_type=pytesseract.Output.DICT)
                
                # Calculate average confidence
                confidences = [int(conf) for conf in data['conf'] if int(conf) > 0]
                if confidences:
                    avg_confidence = sum(confidences) / len(confidences)
                    if avg_confidence > best_confidence:
                        best_confidence = avg_confidence
                        best_result = data
                
                # Also try with binarized image
                _, binary_image = cv2.threshold(processed_image, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
                data_binary = pytesseract.image_to_data(binary_image, config=config,
                                                      output_type=pytesseract.Output.DICT)
                
                confidences_binary = [int(conf) for conf in data_binary['conf'] if int(conf) > 0]
                if confidences_binary:
                    avg_confidence_binary = sum(confidences_binary) / len(confidences_binary)
                    if avg_confidence_binary > best_confidence:
                        best_confidence = avg_confidence_binary
                        best_result = data_binary
                        
            except Exception as e:
                continue
        
        return best_result if best_result else {'text': [], 'conf': [], 'left': [], 'top': [], 'width': [], 'height': []}

    def _try_pattern_matching(self, cell_image):
        """
        Try to recognize simple patterns when OCR fails completely
        """
        try:
            # Check if cell appears to have content based on pixel analysis
            binary = cv2.threshold(cell_image, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)[1]
            
            # Count dark pixels (potential text)
            dark_pixels = np.sum(binary == 0)
            total_pixels = binary.size
            
            if dark_pixels / total_pixels > 0.05:  # At least 5% dark pixels
                # Simple pattern: if there's content, return placeholder
                return "[content]"
            
            return ""
            
        except:
            return ""

    def _extract_cell_text(self, left: int, top: int, right: int, bottom: int) -> str:
        """
        Extract text from specified region with improved OCR handling
        """
        # Add margin to avoid including border lines
        margin = 5
        left = max(0, left + margin)
        top = max(0, top + margin)
        right = min(self.image_width, right - margin)
        bottom = min(self.image_height, bottom - margin)
        
        if right <= left or bottom <= top:
            return ""
        
        cell_image = self.image[top:bottom, left:right]
        
        if cell_image.size == 0:
            return ""
        
        try:
            # Check image statistics first
            mean_intensity = np.mean(cell_image)
            std_intensity = np.std(cell_image)
            
            if std_intensity < 5:  # Very uniform, likely no text
                return ""
            
            min_val, max_val = np.min(cell_image), np.max(cell_image)
            if max_val - min_val < 20:  # Insufficient contrast
                return ""
            
            # Try multiple OCR approaches for this cell
            ocr_data = self._try_multiple_ocr_configs(cell_image)
            
            # Extract and validate text
            valid_texts = []
            for i in range(len(ocr_data.get('text', []))):
                confidence = int(ocr_data['conf'][i]) if i < len(ocr_data['conf']) else 0
                text = ocr_data['text'][i].strip() if i < len(ocr_data['text']) else ""
                
                # More lenient filtering for cell text
                if confidence > 20 and len(text) > 0:
                    clean_text = ''.join(char for char in text if char.isprintable()).strip()
                    if len(clean_text) > 0:
                        # Accept more character types for cell content
                        if (any(char.isalnum() for char in clean_text) or
                            any(char in clean_text for char in ['$', '%', '.', ',', '-', '+', '/', '(', ')', ':', ';'])):
                            valid_texts.append(clean_text)
            
            final_text = ' '.join(valid_texts).strip()
            
            # If no OCR result, try simple template matching for numbers/common patterns
            if not final_text:
                final_text = self._try_pattern_matching(cell_image)
            
            return final_text if len(final_text) > 0 else ""
            
        except Exception as e:
            return ""

    def step3b_refine_lines_with_dl(self) -> Tuple[List[Tuple[int,int,int,int]], List[Tuple[int,int,int,int]]]:
        """
        Step 3b: Refine the initially detected lines using the Deep Learning model.
        This step should be called after step3_detect_table_lines.
        """
        if self.dl_model is None:
            print("DL model not loaded. Skipping line refinement.")
            return self.horizontal_lines, self.vertical_lines

        if self.image_width == 0 or self.image_height == 0:
            print("Image dimensions not set. Cannot refine lines with DL model.")
            return self.horizontal_lines, self.vertical_lines

        initial_lines_combined = []
        for x1, y1, x2, y2 in self.horizontal_lines:
            initial_lines_combined.append([float(x1), float(y1), float(x2), float(y2)])
        for x1, y1, x2, y2 in self.vertical_lines:
            initial_lines_combined.append([float(x1), float(y1), float(x2), float(y2)])

        if not initial_lines_combined:
            print("No initial lines to refine with DL model.")
            return self.horizontal_lines, self.vertical_lines

        lines_np = np.array(initial_lines_combined, dtype=np.float32)

        # Normalize lines to 0-1 range based on image dimensions
        lines_normalized = lines_np.copy()
        lines_normalized[:, [0, 2]] /= self.image_width
        lines_normalized[:, [1, 3]] /= self.image_height
        
        # Ensure no NaN/inf values if image_width/height were zero (though checked above)
        lines_normalized = np.nan_to_num(lines_normalized)

        # Reshape for model: (1, num_lines, 4)
        lines_normalized_batch = np.expand_dims(lines_normalized, axis=0)

        try:
            predicted_lines_normalized_batch = self.dl_model.predict(lines_normalized_batch)
            predicted_lines_normalized = predicted_lines_normalized_batch[0]

            # Denormalize predicted lines
            predicted_lines_denormalized = predicted_lines_normalized.copy()
            predicted_lines_denormalized[:, [0, 2]] *= self.image_width
            predicted_lines_denormalized[:, [1, 3]] *= self.image_height
            
            # Round to nearest integer for pixel coordinates
            predicted_lines_final = np.round(predicted_lines_denormalized).astype(int)

            # Clear old lines and populate with refined ones
            self.horizontal_lines = []
            self.vertical_lines = []
            
            angle_tolerance = 10 # As used in original step3

            for x1_pred, y1_pred, x2_pred, y2_pred in predicted_lines_final:
                # Ensure coordinates are ordered for consistent angle calculation
                # Though the model should ideally output consistent lines.
                # Forcing x1 <= x2 for horizontal, y1 <= y2 for vertical later.

                dx = float(x2_pred - x1_pred)
                dy = float(y2_pred - y1_pred)

                if abs(dx) < 1e-6 and abs(dy) < 1e-6:  # Avoid division by zero for zero-length lines
                    continue

                angle = np.arctan2(dy, dx) * 180 / np.pi
                length = np.sqrt(dx**2 + dy**2)

                if length < 5: # Filter out very short lines post-DL
                    continue

                # Classify and make perfectly horizontal/vertical
                if abs(angle) < angle_tolerance or abs(angle) > (180 - angle_tolerance): # Horizontal
                    avg_y = (y1_pred + y2_pred) // 2
                    self.horizontal_lines.append((min(x1_pred, x2_pred), avg_y, max(x1_pred, x2_pred), avg_y))
                elif abs(angle - 90) < angle_tolerance or abs(angle + 90) < angle_tolerance: # Vertical
                    avg_x = (x1_pred + x2_pred) // 2
                    self.vertical_lines.append((avg_x, min(y1_pred, y2_pred), avg_x, max(y1_pred, y2_pred)))
                # else:
                #     print(f"DL model produced an ambiguous line: ({x1_pred},{y1_pred},{x2_pred},{y2_pred}), angle: {angle}")
        
        except Exception as e:
            print(f"Error during DL model prediction or postprocessing: {e}")

        return self.horizontal_lines, self.vertical_lines

    def step7_generate_docx_table(self, output_path: str) -> str:
        """
        Step 7: Generating docx table with selective borders based on detected lines only
        """
        if not self.table_structure:
            raise ValueError("Please build table structure first")
        
        # Create Word document
        doc = Document()

        # Create table
        rows = len(self.table_structure)
        cols = len(self.table_structure[0]) if rows > 0 else 0

        if rows == 0 or cols == 0:
            # If unable to create table structure, create an empty table
            table = doc.add_table(rows=1, cols=1)
            table.style = 'Table Grid'
            table.rows[0].cells[0].text = "Unable to recognize table content"
        else:
            table = doc.add_table(rows=rows, cols=cols)
            
            # Calculate actual dimensions from detected lines
            h_coords = sorted(set([line[1] for line in self.horizontal_lines]))
            v_coords = sorted(set([line[0] for line in self.vertical_lines]))
            
            # Ensure we have proper coordinates
            if len(h_coords) < 2:
                h_coords = [0, self.image_height]
            if len(v_coords) < 2:
                v_coords = [0, self.image_width]
            
            # Calculate column widths in pixels
            col_widths_px = []
            for j in range(cols):
                left = v_coords[j] if j < len(v_coords) else 0
                right = v_coords[j + 1] if j + 1 < len(v_coords) else self.image_width
                col_widths_px.append(right - left)
            
            # Calculate row heights in pixels
            row_heights_px = []
            for i in range(rows):
                top = h_coords[i] if i < len(h_coords) else 0
                bottom = h_coords[i + 1] if i + 1 < len(h_coords) else self.image_height
                row_heights_px.append(bottom - top)
            
            # Convert pixels to inches (assuming 96 DPI)
            total_width_inches = 6.5  # Standard document width minus margins
            
            # Calculate proportional column widths
            total_width_px = sum(col_widths_px) if sum(col_widths_px) > 0 else 1 # Avoid division by zero
            col_widths_inches = []
            for width_px in col_widths_px:
                width_inches = (width_px / total_width_px) * total_width_inches
                col_widths_inches.append(Inches(max(0.5, width_inches)))  # Minimum 0.5 inches
            
            # Calculate proportional row heights
            max_row_height_inches = 1.0  # Maximum row height
            max_height_px = max(row_heights_px) if row_heights_px else 1
            row_heights_inches = []
            for height_px in row_heights_px:
                height_inches = (height_px / max_height_px) * max_row_height_inches
                row_heights_inches.append(Inches(max(0.25, height_inches)))  # Minimum 0.25 inches
            
            # Apply column widths
            for j, width in enumerate(col_widths_inches):
                if j < len(table.columns):
                    table.columns[j].width = width
            
            # Apply row heights and fill content
            for i, row_data in enumerate(self.table_structure):
                row = table.rows[i]
                
                # Set row height
                if i < len(row_heights_inches):
                    row.height = row_heights_inches[i]
                
                # Fill cell content and apply selective borders
                for j, cell_data in enumerate(row_data):
                    if j < len(row.cells):
                        cell = row.cells[j]
                        cell_text = cell_data['text'].strip() if cell_data['text'].strip() else ""
                        cell.text = cell_text
                        
                        # Set cell properties for better appearance
                        cell.vertical_alignment = 1  # Center vertical alignment
                        
                        # Apply selective borders based on detected lines only
                        self._apply_selective_borders(cell, i, j, h_coords, v_coords)
        
        # Save file
        doc.save(output_path)
        
        return output_path

    def step6_create_table_structure(self) -> List[List]:
        """
        Step 6: Build 2D table data structure (improved version)
        """
        if not self.horizontal_lines and not self.vertical_lines:
            raise ValueError("Not enough table lines detected")
        
        # Get all coordinates and sort
        h_coords = sorted(set([line[1] for line in self.horizontal_lines]))
        v_coords = sorted(set([line[0] for line in self.vertical_lines]))
        
        # Ensure sufficient lines
        if len(h_coords) < 2:
            h_coords = [0, self.image_height]
        if len(v_coords) < 2:
            v_coords = [0, self.image_width]
        
        rows = len(h_coords) - 1
        cols = len(v_coords) - 1
        
        if rows <= 0 or cols <= 0:
            # Final fallback: create 1x1 table
            rows, cols = 1, 1
            h_coords = [0, self.image_height]
            v_coords = [0, self.image_width]
        
        self.table_structure = []
        
        for i in range(rows):
            row = []
            for j in range(cols):
                # Calculate boundaries for each cell
                left = v_coords[j]
                right = v_coords[j + 1] if j + 1 < len(v_coords) else self.image_width
                top = h_coords[i]
                bottom = h_coords[i + 1] if i + 1 < len(h_coords) else self.image_height
                
                # Extract text from cell
                cell_text = self._extract_cell_text(left, top, right, bottom)
                
                row.append({
                    'text': cell_text,
                    'bounds': (left, top, right, bottom)
                })
            self.table_structure.append(row)
        
        return self.table_structure

    def _apply_selective_borders(self, cell, row_idx, col_idx, h_coords, v_coords):
        """
        Apply borders to cell only where lines were actually detected
        """
        from docx.oxml.shared import OxmlElement, qn
        
        # Get cell bounds
        cell_bounds = self.table_structure[row_idx][col_idx]['bounds']
        left, top, right, bottom = cell_bounds
        
        # Tolerance for line matching
        tolerance = 10
        
        # Check which borders should be drawn based on detected lines
        draw_top = False
        draw_bottom = False
        draw_left = False
        draw_right = False
        
        # Check for horizontal lines at cell boundaries
        for x1, y1, x2, y2 in self.horizontal_lines:
            # Check if horizontal line is at top of cell
            if abs(y1 - top) <= tolerance and x1 <= left + tolerance and x2 >= right - tolerance:
                draw_top = True
            # Check if horizontal line is at bottom of cell
            if abs(y1 - bottom) <= tolerance and x1 <= left + tolerance and x2 >= right - tolerance:
                draw_bottom = True
        
        # Check for vertical lines at cell boundaries
        for x1, y1, x2, y2 in self.vertical_lines:
            # Check if vertical line is at left of cell
            if abs(x1 - left) <= tolerance and y1 <= top + tolerance and y2 >= bottom - tolerance:
                draw_left = True
            # Check if vertical line is at right of cell
            if abs(x1 - right) <= tolerance and y1 <= top + tolerance and y2 >= bottom - tolerance:
                draw_right = True
        # Apply borders using docx formatting
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        
        # Create borders element
        tcBorders = tcPr.find(qn('w:tcBorders'))
        if tcBorders is None:
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)
        
        # Helper function to create border element
        def create_border_element(border_name, draw_border):
            # Remove existing border if present
            existing_border = tcBorders.find(qn(f'w:{border_name}'))
            if existing_border is not None:
                tcBorders.remove(existing_border)
            
            # Create new border element
            border_elem = OxmlElement(f'w:{border_name}')
            
            if draw_border:
                # Apply visible border
                border_elem.set(qn('w:val'), 'single')
                border_elem.set(qn('w:sz'), '4')
                border_elem.set(qn('w:space'), '0')
                border_elem.set(qn('w:color'), '000000')
            else:
                # Apply no border
                border_elem.set(qn('w:val'), 'nil')
            
            tcBorders.append(border_elem)
        
        # Apply borders
        create_border_element('top', draw_top)
        create_border_element('bottom', draw_bottom)
        create_border_element('left', draw_left)
        create_border_element('right', draw_right)

    def visualize_detection_results(self, output_path: str = None):
        """
        Visualize detection results (improved version) - only show detected lines, not cell boundaries
        """
        font_scale = 0.7 
        font_thickness = 1 
        font_color = (0, 0, 255)
        text_y_position = 20 
        
        # create original image
        original_labeled = self.original_image.copy()
        cv2.putText(original_labeled, "Original", (10, text_y_position), cv2.FONT_HERSHEY_SIMPLEX, 
                    font_scale, font_color, font_thickness, cv2.LINE_AA)
        # create text_removed image
        text_removed_display = self.text_removed_image.copy()
        text_removed_bgr_labeled = cv2.cvtColor(text_removed_display, cv2.COLOR_GRAY2BGR)
        cv2.putText(text_removed_bgr_labeled, "Text Removed", (10, text_y_position), cv2.FONT_HERSHEY_SIMPLEX, 
                    font_scale, font_color, font_thickness, cv2.LINE_AA)
        # create denoised image
        denoised_bgr_labeled = cv2.cvtColor(self.denoised_image, cv2.COLOR_GRAY2BGR)
        cv2.putText(denoised_bgr_labeled, "Denoised", (10, text_y_position), cv2.FONT_HERSHEY_SIMPLEX, 
                    font_scale, font_color, font_thickness, cv2.LINE_AA)
        # create detected lines image
        detected_image = self.original_image.copy()
        for x1, y1, x2, y2 in self.horizontal_lines: # Draw horizontal lines (red)
            cv2.line(detected_image, (x1, y1), (x2, y2), (0, 0, 255), 2)
        for x1, y1, x2, y2 in self.vertical_lines: # Draw vertical lines (blue)
            cv2.line(detected_image, (x1, y1), (x2, y2), (255, 0, 0), 2)
        cv2.putText(detected_image, "detected lines", (10, text_y_position), cv2.FONT_HERSHEY_SIMPLEX, 
                    font_scale, font_color, font_thickness, cv2.LINE_AA)
        
        visualization_img = np.hstack([original_labeled, text_removed_bgr_labeled, denoised_bgr_labeled, detected_image])
        
        if output_path:
            cv2.imwrite(output_path, visualization_img)
        
        return visualization_img
    
    def process_table_image(self, image_path: str, output_docx_path: str, 
                          visualization_path: str = None) -> Dict:
        """
        Complete table processing workflow
        """
        try:
            # Step 1: Load image
            self.step1_load_image(image_path)

            # Step 2: Remove text
            self.step2_remove_text()

            # Step 2.5: Denoise the image
            self.step2_5_denoise_image()
            
            # Save denoised image
            if self.denoised_image is not None:
                os.makedirs(denoised_folder, exist_ok=True)
                base_name = os.path.splitext(os.path.basename(image_path))[0]
                denoised_path = os.path.join(denoised_folder, f"{base_name}_denoised.jpg")
                cv2.imwrite(denoised_path, self.denoised_image)

            # Step 3: Detect table lines
            self.step3_detect_table_lines()

            # Step 3b: Refine lines with DL model
            if self.dl_model:
                self.step3b_refine_lines_with_dl()

            # Step 4: Merge and filter lines
            self.step4_merge_and_filter_lines()

            # Step 5: Smart line completion
            self.step5_smart_line_completion()

            # Step 6: Create table structure
            self.step6_create_table_structure()

            # Step 7: Generate docx table
            self.step7_generate_docx_table(output_docx_path)
            
            # Generate visualization results
            if visualization_path:
                self.visualize_detection_results(visualization_path)

            # Create processing report
            report = {
                'status': 'success',
                'input_image': image_path,
                'output_docx': output_docx_path,
                'table_size': f"{len(self.table_structure)} x {len(self.table_structure[0]) if self.table_structure else 0}",
                'horizontal_lines': len(self.horizontal_lines),
                'vertical_lines': len(self.vertical_lines),
                'cells_with_text': sum(1 for row in self.table_structure for cell in row if cell['text'].strip())
            }
            
            return report
            
        except Exception as e:
            error_report = {
                'status': 'error',
                'error_message': str(e),
                'input_image': image_path
            }
            return error_report

    def process_folder(self, input_folder: str, output_folder: str = None, 
                      visualization_folder: str = None) -> Dict:
        """
        Process all images in a folder
        """
        if not os.path.exists(input_folder):
            return {
                'status': 'error',
                'error_message': f"Input folder not found: {input_folder}"
            }
        
        # Set default output folders
        if output_folder is None:
            output_folder = os.path.join(input_folder, "output_docx")
        if visualization_folder is None:
            visualization_folder = os.path.join(input_folder, "visualizations")
        
        # Create output directories
        os.makedirs(output_folder, exist_ok=True)
        if visualization_folder:
            os.makedirs(visualization_folder, exist_ok=True)
        
        # Find all image files
        image_extensions = ['.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.tif']
        image_files = []
        
        for file in os.listdir(input_folder):
            if any(file.lower().endswith(ext) for ext in image_extensions):
                image_files.append(file)
        
        if not image_files:
            return {
                'status': 'error',
                'error_message': f"No image files found in {input_folder}"
            }
        
        # Process each image
        results = []
        successful_count = 0
        failed_count = 0
        
        for image_file in image_files:
            image_path = os.path.join(input_folder, image_file)
            
            # Generate output paths
            base_name = os.path.splitext(image_file)[0]
            output_docx = os.path.join(output_folder, f"{base_name}.docx")
            visualization_path = None
            if visualization_folder:
                visualization_path = os.path.join(visualization_folder, f"{base_name}_visualization.jpg")
            
            # Process image
            result = self.process_table_image(
                image_path=image_path,
                output_docx_path=output_docx,
                visualization_path=visualization_path
            )
            
            result['input_file'] = image_file
            results.append(result)
            
            if result['status'] == 'success':
                successful_count += 1
            else:
                failed_count += 1
        
        # Generate summary report
        summary = {
            'status': 'completed',
            'input_folder': input_folder,
            'output_folder': output_folder,
            'visualization_folder': visualization_folder,
            'total_files': len(image_files),
            'successful': successful_count,
            'failed': failed_count,
            'success_rate': f"{(successful_count/len(image_files)*100):.1f}%" if image_files else "0%",
            'results': results
        }
        
        return summary

def main():
    processor = ImprovedTableProcessor()
    
    if os.path.exists(input_folder) and os.path.isdir(input_folder):
        # Batch process entire folder
        result = processor.process_folder(
            input_folder=input_folder,
            output_folder=output_docx_folder,
            visualization_folder=visualization_folder
        )
        
        print("\nBatch Processing Report:")
        print(f"Total files: {result.get('total_files', 0)}")
        print(f"Successful: {result.get('successful', 0)}")
        print(f"Failed: {result.get('failed', 0)}")
        print(f"Success rate: {result.get('success_rate', '0%')}")
        
        # Show individual results summary
        if 'results' in result:
            print("\nIndividual Results:")
            for item in result['results']:
                status = item.get('status', 'unknown')
                filename = item.get('input_file', 'unknown')
                if status == 'success':
                    table_size = item.get('table_size', 'unknown')
                    cells_with_text = item.get('cells_with_text', 0)
                    print(f"✓ {filename}: {table_size} table, {cells_with_text} cells with text")
                else:
                    error_msg = item.get('error_message', 'unknown error')
                    print(f"✗ {filename}: {error_msg}")
        


if __name__ == "__main__":
    main()
